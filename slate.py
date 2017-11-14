from pusher import Pusher

from flask import Flask
from flask import request, send_file, send_from_directory, render_template
from docx import Document
from docx.shared import Pt, Inches

import urllib
import sympy
import json
import os
import os.path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import latex2mathml.converter
from utilities import crossdomain
from StringIO import StringIO

app = Flask(__name__, static_url_path='/static')

pusher = Pusher(
    app_id='428318',
    key='013f281e06025adc67d2',
    secret='30fd86af40447683d9a3',
    cluster='us2',
    ssl=True
)

@app.route('/')
def show_index():

    return render_template('index.html')


@app.route('/messages', methods=['POST', 'OPTIONS'])
@crossdomain(origin='*', headers='Content-Type')
def new_message():
    json = request.json['text']
    time = request.json['time']
    update = request.json['update']
    state_id = request.json['stateID']

    pusher.trigger('messages', 'new_message', {
        'json': json,
        'time': time,
        'update': update,
        'stateID': state_id,
    })

    return "great success!"


@app.route('/create_test_bank', methods=['POST', 'OPTIONS'])
@crossdomain(origin='*', headers='Content-Type')
def create_test_bank():

    slate_json = request.json['questions']

    # Writes to and saves the output doc as output.docx
    return write_to_word_doc(slate_json)

# Method for parsing through different types of values that could be in the dictionaries returned from xml and writes to the word doc
# Right now, can be in 3 forms. Either mixed content, just one content, or an array of content, which is "mixed"
# with different elements but doesn't contain any plain text
# But that resulted in the wrong order
def write_nodes_to_docx(node, table_cell):

    if 'type' in node:
        if node['type'] == 'instructions' or node['type'] == 'choice':
            for child_node in node['nodes']:
                write_nodes_to_docx(child_node, table_cell)

        if node['type'] == 'paragraph':
            # Newlines for blank paragraphs
            if (len(node['nodes']) == 1 and node['nodes'][0]['kind'] == 'text' and
                len(node['nodes'][0]['leaves']) == 1 and
                node['nodes'][0]['leaves'][0]['text'] == ''):
                table_cell.add_paragraph('\n')

            for child_node in node['nodes']:
                write_nodes_to_docx(child_node, table_cell)

        # Run this function again recursively since codeBlock is a block node
        if node['type'] == 'codeblock':
            language = 'plainText' # XML_dict['codeBlock']['@language']
            if language == 'plainText':
                table_cell.paragraphs[0].style = 'Codeblock Plain'
            else:
                table_cell.paragraphs[0].style = 'Codeblock'
            for codeblock_content in node['nodes']:
                write_nodes_to_docx(codeblock_content, table_cell)

        if node['type'] == 'image':
            width = '200'
            img_file_name, _ = urllib.urlretrieve(node['data']['src'], 'tmp.png')
            new_run = table_cell.paragraphs[len(table_cell.paragraphs) - 1].add_run()
            new_run.add_picture(img_file_name, width=Inches(int(width)/75))

        if node['type'] == 'table':
            num_rows = len(node['nodes'])
            num_cols = len(node['nodes'][0]['nodes'])
            table = table_cell.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Zybook Table'
    
            # Delete paragraph in this cell if there isn't any text
            if not len(table_cell.paragraphs[0].text.strip()):
                p = table_cell.paragraphs[0]._element
                p.getparent().remove(p)
                p._p = p._element = None

            for row_idx, row_node in enumerate(node['nodes']):
                for cell_idx, cell_node in enumerate(row_node['nodes']):
                    for node in cell_node['nodes']:
                        write_nodes_to_docx(node, table.rows[row_idx].cells[cell_idx])

    else:
        if node['kind'] == 'text' and 'leaves' in node:
            for leaf in node['leaves']:
                write_nodes_to_docx(leaf, table_cell)

        if node['kind'] == 'leaf':
            if len(node['marks']) and node['marks'][0]['type'] == 'latex':
                table_cell.paragraphs[0].add_run(' ')
                write_latex(node['text'], table_cell)
            else:
                if len(table_cell.paragraphs) <= len(table_cell.tables):
                    table_cell.add_paragraph(node['text'])
                else:
                    table_cell.paragraphs[len(table_cell.paragraphs) - 1].add_run(node['text'])   

def write_latex(latex, cell):
    mml2omml_stylesheet_path = 'mml2omml.xsl'

    # Start by converting latex string to mathml
    mathml_string = ('<math xmlns="http://www.w3.org/1998/Math/MathML">' +
                     latex2mathml.converter.convert(latex.replace('$', '')).decode("utf-8")[6:].replace('amp;', ''))

    # Convert MathML (MML) into Office MathML (OMML) using a XSLT stylesheet
    tree = etree.fromstring(mathml_string)
    xslt = etree.parse(mml2omml_stylesheet_path)

    transform = etree.XSLT(xslt)
    new_dom = transform(tree)

    # table_cell.paragraphs[0]._element.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # new_table = table_cell.add_table(rows=1, cols=1)
    # new_table.allow_autofit = True
    cell.paragraphs[0]._element.append(new_dom.getroot())

# Write to word doc
def write_to_word_doc(slate_json):
    output_file = 'static/testbank.docx'
    template_file = 'template.docx'
    document = Document(template_file)

    questions = slate_json['document']['nodes'][0]['nodes']

    # Section header
    p = document.add_paragraph()
    r = p.add_run('Test bank demo').bold = True

    for question_idx, question_dict in enumerate(questions):
        choices = question_dict['nodes'][1]['nodes']

        prompt_table = document.add_table(rows=1, cols=2)
        prompt_table.style = 'Prompt Table'
        prompt_table.rows[0].cells[0].text = str(question_idx + 1) + ')'
        prompt_table.columns[0].width = Inches(0.4)
        prompt_table.rows[0].cells[0].width = Inches(0.4)
        prompt_table.columns[1].width = Inches(6)
        prompt_table.rows[0].cells[1].width = Inches(6)

        # Table templates are unable to override center valignment in first column, so we'll do it here manually
        tc = prompt_table.rows[0].cells[0]._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), "top")
        tcPr.append(tcVAlign)

        # Question prompt
        write_nodes_to_docx(question_dict['nodes'][0], prompt_table.rows[0].cells[1])

        # Add a new table for the choices
        choice_table = document.add_table(rows=len(choices), cols=2)
        choice_table.style = 'Choice Table'
        choice_table.columns[0].width = Inches(0.5)
        choice_table.columns[1].width = Inches(5.9)

        for cell in choice_table.columns[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), "top")
            tcPr.append(tcVAlign)

        for choice_idx, choice_dict in enumerate(choices):
            choice_table.rows[choice_idx].cells[0].width = Inches(0.5)
            choice_table.rows[choice_idx].cells[1].width = Inches(5.9)

            choice_letter = str(chr(choice_idx + 97))

            r = choice_table.rows[choice_idx].cells[0].paragraphs[0].add_run(choice_letter + ')')
            choice_table.rows[choice_idx].cells[0].paragraphs[0]._element.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if 'correct' in choice_dict['data'] and choice_dict['data']['correct']:
                r.bold = True

            write_nodes_to_docx(choice_dict, choice_table.rows[choice_idx].cells[1])

    document.save(output_file)

    if os.path.isfile('tmp.png'):
        os.remove('tmp.png')

    return 'all done'
