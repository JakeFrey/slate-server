from flask import Flask
from flask import request
from docx import Document
from docx.shared import Pt, Inches

# import urllib.request
import sympy
import json
import os
import os.path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import latex2mathml.converter
from utilities import crossdomain

app = Flask(__name__)

@app.route('/create_test_bank', methods=['POST', 'OPTIONS'])
@crossdomain(origin='*', headers='Content-Type')
def create_test_bank():

    slate_json = request.json['questions']

    # Writes to and saves the output doc as output.docx
    write_to_word_doc(slate_json)

    # will then need to send the file to the client
    return 'Hello, World!'




# Method for parsing through different types of values that could be in the dictionaries returned from xml and writes to the word doc
# Right now, can be in 3 forms. Either mixed content, just one content, or an array of content, which is "mixed"
# with different elements but doesn't contain any plain text
# But that resulted in the wrong order
def xml_to_docx(XML_dict, table_cell):
    if '#' in XML_dict:
        for mixed_content in XML_dict['#']:
            xml_to_docx(mixed_content, table_cell)

    # Run this function again recursively since codeBlock is a complex element
    if 'codeBlock' in XML_dict:
        language = XML_dict['codeBlock']['@language']
        if language == 'plainText':
            table_cell.paragraphs[0].style = 'Codeblock Plain'
        else:
            table_cell.paragraphs[0].style = 'Codeblock'
        for codeblock_content in XML_dict['codeBlock']:
            xml_to_docx({ codeblock_content: XML_dict['codeBlock'][codeblock_content] }, table_cell)
    if 'latex' in XML_dict:
        table_cell.paragraphs[0].add_run(' ')

        # Multiple latex elements, commonly when in table cells or codeblocks
        if len(XML_dict['latex']) > 1:
            for latex in XML_dict['latex']:
                write_latex(latex, table_cell)
        else:
            write_latex(XML_dict['latex'], table_cell)
    if 'image' in XML_dict:
        width = XML_dict['image']['@width']
        #img_file_name, _ = urllib.request.urlretrieve(XML_dict['image']['url']['$'], 'tmp.png')
        #new_run = table_cell.paragraphs[0].add_run()
        #new_run.add_picture(img_file_name, width=Inches(int(width)/75))
    if 'table' in XML_dict:
        num_rows = len(XML_dict['table']['row'])
        num_cols = len(XML_dict['table']['row'][0]['cell'])
        table = table_cell.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Zybook Table'

        # Delete paragraph in this cell if there isn't any text
        if not len(table_cell.paragraphs[0].text.strip()):
            p = table_cell.paragraphs[0]._element
            p.getparent().remove(p)
            p._p = p._element = None

        for row_idx, dict_row in enumerate(XML_dict['table']['row']):
            for cell_idx, dict_cell in enumerate(dict_row['cell']):
                table.rows[row_idx].cells[cell_idx].text = dict_cell['$']    
    if '$' in XML_dict:
        if len(table_cell.paragraphs) <= len(table_cell.tables):
            table_cell.add_paragraph(XML_dict['$'])
        else:
            table_cell.paragraphs[0].add_run(XML_dict['$'])

def write_latex(latex, cell):
    mml2omml_stylesheet_path = 'mml2omml.xsl'

    # Start by converting latex string to mathml
    mathml_string = ('<math xmlns="http://www.w3.org/1998/Math/MathML">' +
                     latex2mathml.converter.convert(str(latex['$'])).decode("utf-8")[6:].replace('amp;', ''))

    # Convert MathML (MML) into Office MathML (OMML) using a XSLT stylesheet
    tree = etree.fromstring(mathml_string)
    xslt = etree.parse(mml2omml_stylesheet_path)

    transform = etree.XSLT(xslt)
    new_dom = transform(tree)

    # table_cell.paragraphs[0]._element.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # new_table = table_cell.add_table(rows=1, cols=1)
    # new_table.allow_autofit = True
    cell.paragraphs[0]._element.append(new_dom.getroot())

# Write to word doc
def write_to_word_doc(slate_json):
    output_file = 'testbank.docx'
    template_file = 'template.docx'

    questions = slate_json['document']['nodes']

    # Convert ordered dictionary to regular dictionary for easier processing
    section_dictionary = json.loads(json.dumps(slate_json))
    document = Document(template_file)

    import pdb; pdb.set_trace()

    # Section header
    p = document.add_paragraph()
    r = p.add_run('Test bank demo').bold = True

    for question_idx, question_dict in enumerate(section_dictionary['testQuestions']['question']):
        prompt_table = document.add_table(rows=1, cols=2)
        prompt_table.style = 'Prompt Table'
        prompt_table.rows[0].cells[0].text = str(question_idx + 1) + ')'
        prompt_table.columns[0].width = Inches(0.4)
        prompt_table.rows[0].cells[0].width = Inches(0.4)
        prompt_table.columns[1].width = Inches(6)
        prompt_table.rows[0].cells[1].width = Inches(6)

        # Table templates are unable to override center valignment in first column, so we'll do it here manually
        tc = prompt_table.rows[0].cells[0]._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), "top")
        tcPr.append(tcVAlign)

        # Question prompt
        xml_to_docx(question_dict['prompt'], prompt_table.rows[0].cells[1])

        # Add a new table for the choices
        choice_table = document.add_table(rows=len(question_dict['choice']), cols=2)
        choice_table.style = 'Choice Table'
        choice_table.columns[0].width = Inches(0.5)
        choice_table.columns[1].width = Inches(5.9)

        for cell in choice_table.columns[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), "top")
            tcPr.append(tcVAlign)

        for choice_idx, choice_dict in enumerate(question_dict['choice']):
            choice_table.rows[choice_idx].cells[0].width = Inches(0.5)
            choice_table.rows[choice_idx].cells[1].width = Inches(5.9)

            choice_letter = str(chr(choice_idx + 97))

            r = choice_table.rows[choice_idx].cells[0].paragraphs[0].add_run(choice_letter + ')')
            choice_table.rows[choice_idx].cells[0].paragraphs[0]._element.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if choice_dict['correct']['$']:
                r.bold = True

            xml_to_docx(choice_dict['label'], choice_table.rows[choice_idx].cells[1])

    document.save(output_file)

    if os.path.isfile('tmp.png'):
        os.remove('tmp.png')
